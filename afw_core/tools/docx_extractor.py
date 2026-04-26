"""Extract text and images from .docx files.

Uses `unstructured` for structured text extraction and `python-docx` for
image extraction with paragraph-level position tracking.
"""

from __future__ import annotations

import base64
import hashlib
import json
import logging
import os
from pathlib import Path
from typing import Annotated

from agent_framework import tool
from docx import Document
from docx.oxml.ns import qn
from unstructured.partition.docx import partition_docx

logger = logging.getLogger(__name__)

DEFAULT_OUTPUT_DIR = os.getenv("DOC_INGEST_OUTPUT_DIR", "output/doc_ingest")


# ---------------------------------------------------------------------------
# Internal helpers (unchanged from the tested CrewAI version)
# ---------------------------------------------------------------------------


def _image_filename(image_bytes: bytes, index: int, content_type: str) -> str:
    """Generate a deterministic filename for an extracted image."""
    digest = hashlib.md5(image_bytes).hexdigest()[:8]
    ext_map = {
        "image/png": "png",
        "image/jpeg": "jpg",
        "image/gif": "gif",
        "image/x-emf": "emf",
        "image/x-wmf": "wmf",
    }
    ext = ext_map.get(content_type, "png")
    return f"img-{index:03d}-{digest}.{ext}"


def _extract_images_with_context(docx_path: Path, images_dir: Path) -> list[dict]:
    """Extract images from a .docx using python-docx, tracking which paragraph they appear in."""
    doc = Document(str(docx_path))
    images_dir.mkdir(parents=True, exist_ok=True)

    image_rels = {}
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image_rels[rel.rId] = rel.target_part

    images = []
    img_index = 0

    for para_idx, para in enumerate(doc.paragraphs):
        blips = para._element.findall(f".//{qn('a:blip')}")
        if not blips:
            continue

        context_parts = []
        for offset in range(max(0, para_idx - 2), para_idx):
            prev_text = doc.paragraphs[offset].text.strip()
            if prev_text:
                context_parts.append(prev_text)
        if para.text.strip():
            context_parts.append(para.text.strip())

        for blip in blips:
            r_embed = blip.get(qn("r:embed"))
            if not r_embed or r_embed not in image_rels:
                continue

            part = image_rels[r_embed]
            image_bytes = part.blob
            content_type = part.content_type

            fname = _image_filename(image_bytes, img_index, content_type)
            img_path = images_dir / fname
            img_path.write_bytes(image_bytes)
            img_index += 1

            images.append({
                "path": str(img_path),
                "context": " | ".join(context_parts) if context_parts else "",
                "paragraph_index": para_idx,
            })

    return images


def _extract_tables_as_markdown(docx_path: Path) -> list[str]:
    """Extract all tables from a .docx as Markdown-formatted strings."""
    doc = Document(str(docx_path))
    tables_md = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows.append(cells)
        if not rows:
            continue
        header = "| " + " | ".join(rows[0]) + " |"
        separator = "| " + " | ".join("---" for _ in rows[0]) + " |"
        body_lines = []
        for row in rows[1:]:
            while len(row) < len(rows[0]):
                row.append("")
            body_lines.append("| " + " | ".join(row[: len(rows[0])]) + " |")
        md = header + "\n" + separator
        if body_lines:
            md += "\n" + "\n".join(body_lines)
        tables_md.append(md)
    return tables_md


def _extract_docx(docx_path: Path, output_dir: Path) -> dict:
    """Extract all elements from a single .docx file."""
    images_dir = output_dir / "images"

    elements = partition_docx(filename=str(docx_path))
    tables_md = _extract_tables_as_markdown(docx_path)
    table_index = 0

    text_segments: list[dict[str, str]] = []
    for el in elements:
        el_type = type(el).__name__
        text = el.text.strip() if hasattr(el, "text") and el.text else ""
        if text:
            if el_type == "Table" and table_index < len(tables_md):
                text_segments.append({"type": "Table", "text": tables_md[table_index]})
                table_index += 1
            else:
                text_segments.append({"type": el_type, "text": text})

    images_raw = _extract_images_with_context(docx_path, images_dir)

    images_raw.sort(key=lambda im: im["paragraph_index"])
    insert_offset = 0
    for img in images_raw:
        ctx = img.get("context", "")
        inserted = False
        if ctx:
            ctx_parts = [p.strip() for p in ctx.split(" | ") if p.strip()]
            for idx in range(insert_offset, len(text_segments)):
                seg_text = text_segments[idx]["text"]
                for ctx_part in ctx_parts:
                    if ctx_part and ctx_part in seg_text:
                        placeholder = {
                            "type": "image",
                            "text": f"[IMAGE: {Path(img['path']).name}]",
                        }
                        text_segments.insert(idx + 1, placeholder)
                        insert_offset = idx + 2
                        inserted = True
                        break
                if inserted:
                    break
        if not inserted:
            text_segments.append({
                "type": "image",
                "text": f"[IMAGE: {Path(img['path']).name}]",
            })
            insert_offset = len(text_segments)

    images = [{"path": img["path"], "context": img["context"]} for img in images_raw]
    return {"text_segments": text_segments, "images": images}


def _extract_all_docx(input_dir: Path, output_dir: Path) -> list[dict]:
    """Extract from all .docx files in a directory."""
    results = []
    for docx_file in sorted(input_dir.glob("*.docx")):
        logger.info("Extracting: %s", docx_file.name)
        result = _extract_docx(docx_file, output_dir)
        result["source_file"] = docx_file.name
        results.append(result)
    return results


# ---------------------------------------------------------------------------
# MAF tool
# ---------------------------------------------------------------------------


@tool
def extract_docx_files(
    input_dir: Annotated[str, "Path to the folder containing .docx files"],
    output_dir: Annotated[str, "Path to save extracted images"] = DEFAULT_OUTPUT_DIR,
) -> str:
    """Extract text and images from all .docx files in a directory.

    Returns a JSON array of extraction results per file, each with
    'source_file', 'text_segments', and 'images'.
    """
    results = _extract_all_docx(Path(input_dir), Path(output_dir))
    return json.dumps(results, ensure_ascii=False)
